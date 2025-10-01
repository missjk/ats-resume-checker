
import re
import PyPDF2
import docx
import mammoth
from datetime import datetime
import os
from transformers import AutoTokenizer, AutoModelForTokenClassification
import torch
from config import Config
from criteria_evaluator import CriteriaEvaluator 

class ResumeParser:
    def __init__(self):
        # Define comprehensive keyword sets for different criteria
        # --- BERT model and tokenizer initialization (Added) ---

        self.tokenizer = AutoTokenizer.from_pretrained("dslim/bert-base-NER")  # Added
        self.model = AutoModelForTokenClassification.from_pretrained("dslim/bert-base-NER")  # Added
        self.label_list = self.model.config.id2label  # Added

        # --- BERT model accuracy (for UI display) ---
        self.bert_model_accuracy = getattr(self.model.config, "id2label", None)
        self.bert_model_accuracy_score = getattr(self.model.config, "num_labels", None)
        # --------------------------------------------------------
        
        self.company_law_keywords = [
            "company law", "corporate law", "companies act", "corporate governance",
            "mergers and acquisitions", "m&a", "corporate compliance", "board meetings",
            "shareholder rights", "corporate restructuring", "due diligence",
            "corporate finance", "securities law", "listing regulations"
        ]

        self.contract_law_keywords = [
            "contract law", "contracts", "contractual", "agreement law",
            "breach of contract", "contract drafting", "commercial contracts",
            "contract negotiation", "terms and conditions", "contract management",
            "specific performance", "damages", "remedies", "consideration"
        ]

        self.legal_research_keywords = [
            "legal research", "research paper", "legal writing", "case study",
            "legal analysis", "jurisprudence", "legal opinion", "case brief",
            "legal memorandum", "research methodology", "legal database",
            "westlaw", "manupatra", "lexis", "legal precedent", "research work"
        ]

        self.moot_court_keywords = [
            "moot court", "mooting", "moot competition", "appellate advocacy",
            "oral arguments", "brief writing", "advocacy", "philip jessup",
            "vis moot", "national moot", "international moot", "arbitration moot",
            "constitutional law moot", "criminal law moot", "moot elimination",
            "mediation competition", "dispute resolution competition"
        ]

        self.tier_firms = Config.get_tier_firms()

    # def process_text_with_bert(self, text):
    #     inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
    #     outputs = self.model(**inputs)
    #     predictions = torch.argmax(outputs.logits, dim=2)

    #     tokens = self.tokenizer.convert_ids_to_tokens(inputs["input_ids"][0])
    #     entities = []
    #     for token, prediction in zip(tokens, predictions[0]):
    #         label = self.label_list[prediction.item()]
    #         if label != "O":
    #             entities.append((token, label))
    #     return entities
    def process_text_with_bert(self, text):
        inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
        outputs = self.model(**inputs)
        
        # Get predictions and confidence scores
        predictions = torch.argmax(outputs.logits, dim=2)
        
        # Calculate confidence scores using softmax
        probabilities = torch.softmax(outputs.logits, dim=2)
        confidence_scores = torch.max(probabilities, dim=2)[0]  # Get max probability for each token
        
        tokens = self.tokenizer.convert_ids_to_tokens(inputs["input_ids"][0])
        entities = []
        confidence_sum = 0.0
        valid_predictions = 0
        
        for token, prediction, confidence in zip(tokens, predictions[0], confidence_scores[0]):
            label = self.label_list[prediction.item()]
            if label != "O":  # Only count non-"Other" predictions
                entities.append((token, label))
                confidence_sum += confidence.item()
                valid_predictions += 1
        
        # Calculate average confidence for this text
        avg_confidence = (confidence_sum / valid_predictions * 100) if valid_predictions > 0 else 0.0
        
        return entities, round(avg_confidence, 2)



    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file using PyPDF2"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"Error reading page {page_num}: {e}")
                        continue
                return text
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    # def extract_text_from_docx(self, file_path):
    #     """Extract text from DOCX file"""
    #     try:
    #         doc = docx.Document(file_path)
    #         text = ""
    #         for paragraph in doc.paragraphs:
    #             text += paragraph.text + "\n"

    #         # Also extract text from tables
    #         for table in doc.tables:
    #             for row in table.rows:
    #                 for cell in row.cells:
    #                     text += cell.text + "\t"
    #                 text += "\n"
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            with open(file_path, "rb") as docx_file:
                import mammoth
                result = mammoth.extract_raw_text(docx_file)
                text = result.value

            if len(text) < 50:
                return "Insufficient text content"
            return text

        except Exception as e:
            return "Error extracting DOCX: " + str(e)


    def extract_cgpa(self, text):
        """Extract CGPA/GPA from resume text with IMPROVED pattern matching"""
        # Clean text for better pattern matching
        text = re.sub(r'\s+', ' ', text)

        # Enhanced patterns to catch ALL CGPA/GPA/Percentage formats
        patterns = [
            # Standard formats: CGPA: 8.5, GPA: 3.8, G.P.A: 7.5
            r'(?:CGPA|GPA|G\.P\.A|C\.G\.P\.A|Cumulative GPA)\s*[:-–]?\s*(?:\(Avg\.\)\s*[-–]\s*)?([0-9]+(?:\.[0-9]+)?)(?:\s*/\s*([0-9]+))?',
            # Reverse formats: 8.5 CGPA, 3.8 GPA  
            r'([0-9]+(?:\.[0-9]+)?)\s*(?:CGPA|GPA|G\.P\.A|C\.G\.P\.A)(?:\s*/\s*([0-9]+))?',
            # Fraction formats: 8.5/10, 3.8/4.0
            r'(?:CGPA|GPA|G\.P\.A)\s*[:-–]?\s*([0-9]+(?:\.[0-9]+)?)\s*/\s*([0-9]+(?:\.[0-9]+)?)',
            # Academic performance indicators
            r'(?:Academic Performance|Overall Grade|Cumulative Grade)\s*[:-–]?\s*([0-9]+(?:\.[0-9]+)?)(?:\s*/\s*([0-9]+))?',
            # Percentage formats: 66% (1st Class), 95.4%
            r'\b([0-9]+(?:\.[0-9]+)?)\s*%\s*(?:\([^)]*\))?',
            # B.A. LL.B. (Hons.) – School Name – 66% format
            r'[–-]\s*([0-9]+(?:\.[0-9]+)?)\s*%\s*(?:\([^)]*\))?'
        ]

        for i, pattern in enumerate(patterns):
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    if i >= 4:  # Percentage patterns (last two)
                        percentage = float(match.group(1))
                        # Convert percentage to CGPA estimate
                        if percentage >= 90:
                            cgpa = 9.0 + (percentage - 90) / 10
                        elif percentage >= 80:
                            cgpa = 8.0 + (percentage - 80) / 10
                        elif percentage >= 70:
                            cgpa = 7.0 + (percentage - 70) / 10
                        elif percentage >= 60:
                            cgpa = 6.0 + (percentage - 60) / 10
                        else:
                            cgpa = percentage / 10  # Basic conversion
                        return min(cgpa, 10.0)
                    else:
                        cgpa = float(match.group(1))
                        max_cgpa = float(match.group(2)) if len(match.groups()) > 1 and match.group(2) else None

                        # Normalize to 10 scale if necessary
                        if max_cgpa:
                            if max_cgpa == 4.0:
                                cgpa = (cgpa / 4.0) * 10.0
                            elif max_cgpa == 5.0:
                                cgpa = (cgpa / 5.0) * 10.0

                        # Validate CGPA range
                        if 0 <= cgpa <= 10:
                            return cgpa
                except ValueError:
                    continue

        return None

    def extract_academic_year(self, text):
        """Extract current academic year from resume with COMPLETE implementation"""
        current_year = datetime.now().year

        # Enhanced patterns to identify academic year
        year_patterns = [
            # Direct year mentions: "3rd year", "fourth year"
            r'\b([1-5])(?:st|nd|rd|th)\s+year\b',
            r'\b(?:first|second|third|fourth|fifth|1st|2nd|3rd|4th|5th)\s+year\b',

            # Semester patterns: "6th semester", "semester 8"
            r'\b([1-9]|10)(?:st|nd|rd|th)?\s+semester\b',
            r'\bsemester\s+([1-9]|10)\b',

            # Year level: "Year 3", "Level 4"
            r'\b(?:year|level)\s+([1-5])\b',

            # Academic year ranges: "2021-2026", "2024-25", "(2021-26)"
            r'\(?\s*(20[2-9][0-9])\s*[-–]\s*(?:20)?([2-9][0-9])\s*\)?',

            # Expected graduation: "graduating in 2025", "class of 2024"
            r'\b(?:graduating|graduation|class of)\s+(?:in\s+)?(20[2-9][0-9])\b',

            # Current enrollment: "currently in 3rd year"
            r'\bcurrently\s+(?:in\s+)?(?:([1-5])(?:st|nd|rd|th)|([1-9])(?:st|nd|rd|th)?)\s+(?:year|semester)\b',

            # Final-year patterns: "Final-year B.A. LL.B."
            r'\bfinal[-\s]?year\b'
        ]

        year_words_to_numbers = {
            'first': 1, 'second': 2, 'third': 3, 'fourth': 4, 'fifth': 5,
            '1st': 1, '2nd': 2, '3rd': 3, '4th': 4, '5th': 5
        }

        for pattern_index, pattern in enumerate(year_patterns):
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    # Handle final-year pattern
                    if pattern_index == 8:  # final-year pattern
                        return 5

                    # Handle semester patterns
                    elif pattern_index in [2, 3]:  # semester patterns
                        if match.group(1) and match.group(1).isdigit():
                            semester_number = int(match.group(1))
                            year_number = (semester_number + 1) // 2
                            if 1 <= year_number <= 5:
                                return year_number

                    # Handle academic year ranges like "2021-26" or "(2021–2026)"
                    elif pattern_index == 5:  # Academic year ranges
                        start_year = int(match.group(1))
                        end_year_str = match.group(2)

                        # Handle 2-digit or 4-digit end year
                        if len(end_year_str) == 2:
                            if int(end_year_str) < 50:
                                end_year = 2000 + int(end_year_str)
                            else:
                                end_year = 1900 + int(end_year_str)
                        else:
                            end_year = int(end_year_str)

                        # Calculate current year in the program
                        years_in_college = current_year - start_year + 1

                        # Validate range (1-5 years for law programs)
                        if 1 <= years_in_college <= 6:  # Allow up to 6 for flexibility
                            return min(years_in_college, 5)

                    # Handle graduation year patterns
                    elif pattern_index == 6:  # graduation patterns
                        grad_year = int(match.group(1))
                        years_remaining = grad_year - current_year
                        if years_remaining == 1:
                            return 5
                        elif years_remaining == 2:
                            return 4
                        elif years_remaining == 3:
                            return 3

                    # Handle direct year mentions
                    else:
                        year_text = match.group(1) if match.group(1) else match.group(0).split()[0]
                        if year_text in year_words_to_numbers:
                            year_number = year_words_to_numbers[year_text]
                        elif year_text.isdigit():
                            year_number = int(year_text)
                        else:
                            continue

                        if 1 <= year_number <= 5:
                            return year_number

                except (ValueError, IndexError, AttributeError):
                    continue

        return None

    def check_course_keywords(self, text, keyword_list):
        """Check if any course-related keywords are present in text"""
        text_lower = text.lower()
        found_keywords = []

        for keyword in keyword_list:
            keyword_lower = keyword.lower()
            if keyword_lower in text_lower:
                found_keywords.append(keyword)

        return len(found_keywords) > 0, found_keywords

    def extract_experience(self, text):
        """Extract comprehensive legal experience information"""
        experience_info = {
            'legal_research': False,
            'moot_court': False,
            'internships': [],
            'publications': [],
            'tier_firm_internship': False,
            'ma_moot_experience': False,
            'faculty_recommendation': False,
            'legalogic_previous': False
        }

        # Check for legal research experience
        legal_research_found, _ = self.check_course_keywords(text, self.legal_research_keywords)
        experience_info['legal_research'] = legal_research_found

        # Check for moot court experience
        moot_court_found, _ = self.check_course_keywords(text, self.moot_court_keywords)
        experience_info['moot_court'] = moot_court_found

        # Check for M&A specific moot experience
        ma_moot_patterns = [
            r'\b(?:m&a|merger|acquisition)\s+moot\b',
            r'\bcorporate\s+law\s+moot\b',
            r'\bcompany\s+law\s+moot\b'
        ]
        for pattern in ma_moot_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                experience_info['ma_moot_experience'] = True
                break

        # Extract internship information and check for tier firms
        for firm in self.tier_firms:
            if firm.lower() in text.lower():
                experience_info['tier_firm_internship'] = True
                experience_info['internships'].append(f"Experience at {firm}")
                break

        return experience_info

    def parse_resume(self, file_path):
        
        """Main function to parse resume and extract all relevant information"""
        if not os.path.exists(file_path):
            return {"error": "File not found"}

        file_extension = file_path.lower().split('.')[-1]

        if file_extension == 'pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == 'docx':
            text = self.extract_text_from_docx(file_path)
        else:
            return {"error": "Unsupported file format"}

        if not text or text.startswith("Error"):
            return {"error": f"Could not extract text from file: {text}"}

        if len(text.strip()) < 50:
            return {"error": "Insufficient text content in resume"}

        # Debug: Print extracted text for troubleshooting
        print(f"\n=== DEBUG: Parsing {os.path.basename(file_path)} ===")
        print(f"Text preview: {text[:200]}...")

        
        # --- BERT processing step added ---
        entities, bert_confidence = self.process_text_with_bert(text)
        print(f"BERT extracted entities: {entities}")
        print(f"BERT confidence score: {bert_confidence}%")
        # ------------------------------------

        # Extract all information with debug output
        cgpa = self.extract_cgpa(text)
        academic_year = self.extract_academic_year(text)
        company_law = self.check_course_keywords(text, self.company_law_keywords)[0]
        contract_law = self.check_course_keywords(text, self.contract_law_keywords)[0]
        experience = self.extract_experience(text)

        print(f"CGPA extracted: {cgpa}")
        print(f"Academic year: {academic_year}")
        print(f"Company law: {company_law}")
        print(f"Contract law: {contract_law}")
        print(f"=== END DEBUG ===\n")
        
        parsed_info = {
            'filename': os.path.basename(file_path),
            'cgpa': cgpa,
            'academic_year': academic_year,
            'company_law': company_law,
            'contract_law': contract_law,
            'experience': experience,
            'text_length': len(text),
            'raw_text_preview': text[:Config.TEXT_PREVIEW_LENGTH] + "..." if len(text) > Config.TEXT_PREVIEW_LENGTH else text,
            'bert_confidence': bert_confidence, # Include BERT confidence score
            'bert_entities_count': len(entities)



        }
        preference_score, preference_details = CriteriaEvaluator().calculate_preference_score(parsed_info)
        parsed_info['preference_details'] = preference_details  # Ensure preference_details is included

        parsed_info['preference'] = preference_score
        parsed_info['preference_explanation'] = "; ".join(
            [f"{k.replace('_', ' ').title()}: {v} points" for k, v in preference_details.items() if v > 0]
        )
        # Debug: Print preference details to verify
        print(f"Preference details: {preference_details}")
        return parsed_info

